"""One-off: build docs/breakage-policy.docx from the policy content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Breakage Calculation Policy", level=0)

meta = doc.add_paragraph()
meta.add_run("Status: ").bold = True
meta.add_run("Internal reference\n")
meta.add_run("Last updated: ").bold = True
meta.add_run("12 June 2026\n")
meta.add_run("Owner: ").bold = True
meta.add_run("Product (with inputs from Finance)")

doc.add_heading("1. Background: how a redemption works", level=1)
doc.add_paragraph(
    "When a user redeems points against a gift card, the platform does not "
    "immediately purchase the gift card from the vendor. Instead:"
)
for step in [
    "The user places a redemption order using their points.",
    "The platform sends the user a click-to-activate link (via email).",
    "Only when the user clicks the link do we call the vendor API — this is the "
    "moment the gift card is actually issued and paid for. This is when the actual "
    "redemption happens.",
]:
    doc.add_paragraph(step, style="List Number")
doc.add_paragraph(
    "A meaningful share of these links are never clicked. Points that are deducted "
    "from the user but never result in a vendor purchase — either because the link "
    "was never activated, or because no redemption was ever initiated at all — are "
    "what we call breakage."
)

doc.add_heading("2. What counts as breakage", level=1)
doc.add_paragraph("There are two distinct sources of breakage:")
t = doc.add_table(rows=3, cols=2)
t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "Source"
t.rows[0].cells[1].text = "Description"
t.rows[1].cells[0].text = "Unactivated links"
t.rows[1].cells[1].text = (
    "A redemption order was placed and a click-to-activate link was sent, but the "
    "user never clicked it."
)
t.rows[2].cells[0].text = "Dormant points"
t.rows[2].cells[1].text = (
    "Points sitting on the platform that were never redeemed at all — no link was "
    "ever generated."
)
doc.add_paragraph(
    "Finance applies different recognition rules to each, and the rules also differ "
    "by currency."
)

doc.add_heading("3. INR redemptions (click-to-activate links)", level=1)
inr_points = [
    ("Link validity: ",
     "INR click-to-activate links are valid for 6 months from issuance. After 6 "
     "months, the link expires and the redemption is no longer active."),
    ("Revalidation: ",
     "Even after expiry, if a user contacts us and asks for the link to be "
     "revalidated, we do revalidate it. Expiry is therefore a soft cutoff from the "
     "user's perspective."),
    ("Breakage recognition: ",
     "Finance calculates INR link breakage annually in March, at the close of the "
     "financial year (April–March). All links generated in the previous fiscal year "
     "that expired without being activated are counted as breakage, at their "
     "monetary value."),
    ("Post-recognition adjustment: ",
     "If a link that was already counted as breakage is later revalidated and "
     "activated, the breakage figure is adjusted/reversed — the activation is "
     "netted off against breakage."),
]
for lead, body in inr_points:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(lead).bold = True
    p.add_run(body)

doc.add_heading("4. Non-INR redemptions (click-to-activate links)", level=1)
for lead, body in [
    ("Link validity: ",
     "For all currencies other than INR, there is no expiry — links remain active "
     "indefinitely."),
    ("Breakage recognition: ",
     "A non-INR link is counted as breakage only if it remains unactivated for 2 "
     "years. Until the 2-year mark, an inactive link is not breakage — it is simply "
     "a pending redemption."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(lead).bold = True
    p.add_run(body)

doc.add_heading("Exception: the March 2026 calculation", level=2)
doc.add_paragraph(
    "For a period, the redemption emails for USD and other non-INR denominations "
    "stated that the links would be valid for 6 months only. On the basis of that "
    "communicated validity, a subset of those non-INR links was treated under the "
    "6-month rule and included as breakage in the March 2026 calculation, rather "
    "than waiting for the standard 2-year window. This is a known one-time "
    "deviation from the standard non-INR policy and should be kept in mind when "
    "comparing breakage figures across years."
)

doc.add_heading("5. Dormant points (no link ever generated)", level=1)
doc.add_paragraph(
    "Separately from link-based breakage, Finance also reviews points that have "
    "never been redeemed at all:"
)
for body in [
    "Points that remain dormant on the platform for 2 years — with no redemption "
    "initiated and no link generated — are identified by Finance.",
    "The monetary value of these points is computed and counted as breakage after "
    "the 2-year mark.",
]:
    doc.add_paragraph(body, style="List Bullet")
doc.add_paragraph("This applies across all currencies.")

doc.add_heading("6. Summary of recognition rules", level=1)
rows = [
    ("Category", "Validity / window", "When counted as breakage"),
    ("INR link, unactivated",
     "Link valid 6 months (revalidation on request)",
     "At fiscal year close (March), for links issued in the previous FY that "
     "expired unactivated; reversed if later revalidated and activated"),
    ("Non-INR link, unactivated", "No expiry", "After 2 years of inactivity"),
    ("Non-INR link (6-month-validity email cohort)",
     "6 months, as communicated in email",
     "One-time inclusion in the March 2026 calculation"),
    ("Dormant points (no link)", "n/a",
     "After 2 years of dormancy, at monetary value"),
]
t = doc.add_table(rows=len(rows), cols=3)
t.style = "Light Grid Accent 1"
for i, row in enumerate(rows):
    for j, cell in enumerate(row):
        t.rows[i].cells[j].text = cell

doc.add_heading("7. Relation to the Breakage Tool", level=1)
doc.add_paragraph(
    "The Breakage Tool operates on link-level data and flags a record as breakage "
    "when Final Redemption Status == 0 (i.e., the link was never activated), "
    "regardless of how long the link has been outstanding. It uses a statistical "
    "maturity threshold to avoid overstating breakage for recent cohorts, which is "
    "an analytical convention rather than the accounting one."
)
doc.add_paragraph(
    "The figures in the tool will therefore not match Finance's recognized breakage "
    "one-for-one: Finance recognizes breakage on the fixed timelines above (FY "
    "close for INR, 2 years for non-INR and dormant points), while the tool reports "
    "point-in-time unactivated status for trend and cohort analysis. Both views are "
    "correct for their purpose; this document describes the accounting view."
)

doc.save("docs/breakage-policy.docx")
print("Saved docs/breakage-policy.docx")
